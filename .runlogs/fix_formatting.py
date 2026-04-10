# -*- coding: utf-8 -*-
"""Align new project run-level formatting with 慧云智学."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import docx
from docx.shared import Pt, RGBColor

RESUME = r'C:\Users\35502\Desktop\简历.docx'

FONT_NAME = 'Arial'
TITLE_SIZE = Pt(11)
DATE_COLOR = RGBColor(0x44, 0x44, 0x44)


def clear_runs(p):
    """Remove all existing runs from a paragraph."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def add_run(p, text, *, bold=False, size=None, color=None):
    r = p.add_run(text)
    r.font.name = FONT_NAME
    # 中文字体回退 —— 设置 eastAsia 让中文也走 Arial 的 CJK fallback（和原段落一致）
    from docx.oxml.ns import qn
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    if bold:
        r.bold = True
    if size:
        r.font.size = size
    if color is not None:
        r.font.color.rgb = color
    return r


def main():
    doc = docx.Document(RESUME)

    # —— [25] 项目标题行 ——
    p = doc.paragraphs[25]
    assert '智策' in p.text
    clear_runs(p)
    add_run(p, '智策 · AI 决策助手', bold=True, size=TITLE_SIZE)
    add_run(p, '\t2026-02 至 2026-04', color=DATE_COLOR)

    # —— [26]/[27]/[28] 标签：正文 ——
    LABELED = [
        (26, '技术栈：', '技术栈：SpringBoot + SpringCloudAlibaba(Nacos) + Spring AI Alibaba + Milvus + Redis(Redisson) + MySQL + MyBatis-Plus + RabbitMQ + MCP协议 + 通义千问-Max 等'),
        (27, '项目介绍：', '项目介绍：面向企业内部场景的 AI 决策助手平台，集成通义千问-Max/Deepseek-V3 大模型，提供基于 RAG 的企业知识问答、多数据源 Tool Calling（MySQL/Redis/外部API）、客服工单自动化以及 MCP 协议动态工具扩展能力。前后端分离，支持流式对话、多轮上下文记忆与会话隔离。'),
        (28, '个人职责：', '个人职责：负责后端核心链路，包括 RAG 检索管线、ChatMemory 持久化、Tool Calling 工具体系、MCP 动态客户端治理及文档摄入管线的设计与实现。'),
    ]
    for idx, label, full in LABELED:
        p = doc.paragraphs[idx]
        assert p.text.startswith(label), f'paragraph[{idx}] 不以 "{label}" 开头'
        body = full[len(label):]
        clear_runs(p)
        add_run(p, label, bold=True)
        add_run(p, body)

    # —— [29..34] bullet 段落：单 run，Arial ——
    BULLETS = {
        29: '基于 Milvus 落地 RAG 混合检索方案：稠密向量（text-embedding-v3, 1024 维）结合 BM25 稀疏向量双路召回，通过 RRF（Reciprocal Rank Fusion）算法融合排序，相比单一语义检索显著提升长尾查询命中率与召回精度。',
        30: '基于 Apache Tika + TokenTextSplitter 构建文档摄入管线，按 token 粒度切片并注入 kb_code/doc_id/chunk_index 元数据实现多知识库隔离与来源溯源；针对 Milvus 写入区分"永久性故障"（解析失败直接标记 FAILED）与"瞬时故障"（网络超时抛出交由 RabbitMQ 自动重试），保障最终一致性。',
        31: '基于 Redisson 自研 ChatMemoryRepository 实现对话上下文持久化：Redis 承载热数据保障低延迟响应，saveAll 时同步发布 MQ 消息异步写入 MySQL 归档，并通过 Redis pending 集合 + 定时任务兜底，避免消息丢失导致的数据不一致。',
        32: '基于 CompletableFuture + 自定义线程池实现 Tool Calling 并行执行：对 LLM 返回的多个工具调用 fan-out 并发触发，单调用降级为同步路径避免线程池开销；结合 ToolCatalog 动态合并本地 FunctionToolCallback 与 MCP 远端工具，支持大模型按意图编排 MySQL/Redis/外部API/知识库/工单等多源能力，显著降低多工具链路的端到端延迟。',
        33: '自研 McpToolRegistry 解决 MCP Server 动态接入问题：通过 initialized=false 实现启动期非阻塞；@Scheduled 后台定时探活，接入 ping() 作 liveness 检测，失败后 closeGracefully 并在下一轮自动重建 SSE 连接实现故障自愈；ToolCatalog 通过 volatile 缓存 + ObjectProvider 动态合并本地与 MCP 工具，无需重启即可热挂载新工具。',
        34: '基于策略模式封装工单类型与处理人映射规则，通过 Tool Calling 暴露"创建/查询/更新/关闭"四类操作让 LLM 直接驱动工单流转，并联动 RabbitMQ 异步触发邮件通知，实现"AI 对话 → 工单落库 → 处理人触达"一站式闭环。',
    }
    for idx, text in BULLETS.items():
        p = doc.paragraphs[idx]
        clear_runs(p)
        add_run(p, text)

    doc.save(RESUME)
    print('OK — run-level formatting realigned')


if __name__ == '__main__':
    main()
