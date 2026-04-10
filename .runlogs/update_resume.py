# -*- coding: utf-8 -*-
"""Rename project, replace bullet 4, and refresh 个人技能."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import docx

RESUME = r'C:\Users\35502\Desktop\简历.docx'

# 精确文本替换：(原文开头前缀, 新全文)
REPLACEMENTS = [
    # —— 个人技能：AI 应用开发 ——
    (
        'AI应用开发: 熟练运用Spring AI Alibaba框架',
        'AI应用开发: 熟练运用 Spring AI Alibaba 框架，集成通义千问-Max、Deepseek-V3 大模型；'
        '掌握 Prompt 工程优化技巧；深入理解 RAG 原理，基于 Milvus 向量库落地稠密向量 + BM25 稀疏向量 + RRF 融合的混合检索方案，'
        '熟悉文档摄入管线（Tika 解析 + Token 切片 + 元数据注入）；具备 AI 应用后端快速开发能力。',
    ),
    # —— 个人技能：AI 系统工程化 ——
    (
        'AI系统工程化: 熟悉大模型与后端系统集成方式',
        'AI系统工程化: 熟悉大模型与后端系统集成方式；掌握 Tool Calling 核心技术并基于 CompletableFuture 实现多工具并行调用；'
        '了解 MCP 协议与动态客户端生命周期治理（启动非阻塞、ping 探活、SSE 自愈重连）；'
        '通过 Redisson 自研 ChatMemoryRepository 实现对话上下文持久化；具备 AI 服务稳定性保障设计能力。',
    ),
    # —— 项目名重命名 ——
    (
        '企业智能决策助手',
        '智策 · AI 决策助手\t2026-02 至 2026-04',
    ),
    # —— bullet #4：替换为"并行 Tool Calling 执行" ——
    (
        '基于 Spring AI FunctionToolCallback 封装 MySQL 查询',
        '基于 CompletableFuture + 自定义线程池实现 Tool Calling 并行执行：对 LLM 返回的多个工具调用 fan-out 并发触发，'
        '单调用降级为同步路径避免线程池开销；结合 ToolCatalog 动态合并本地 FunctionToolCallback 与 MCP 远端工具，'
        '支持大模型按意图编排 MySQL/Redis/外部API/知识库/工单等多源能力，显著降低多工具链路的端到端延迟。',
    ),
]


def main():
    doc = docx.Document(RESUME)

    hits = {prefix: False for prefix, _ in REPLACEMENTS}
    for p in doc.paragraphs:
        for prefix, new_text in REPLACEMENTS:
            if hits[prefix]:
                continue
            if p.text.startswith(prefix):
                # 清空所有 run，保留第一个 run 的格式信息并写入新文本
                if p.runs:
                    first_run = p.runs[0]
                    first_run.text = new_text
                    # 删掉其余 run 避免残留
                    for r in p.runs[1:]:
                        r.text = ''
                else:
                    p.add_run(new_text)
                hits[prefix] = True
                break

    missing = [k for k, v in hits.items() if not v]
    if missing:
        raise RuntimeError(f'未命中以下前缀：{missing}')

    doc.save(RESUME)
    print('OK — all', len(REPLACEMENTS), 'replacements applied')


if __name__ == '__main__':
    main()
