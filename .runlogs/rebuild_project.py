# -*- coding: utf-8 -*-
"""Rebuild new project by deep-copying 慧云智学 paragraphs."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
import docx
from docx.oxml.ns import qn

RESUME = r'C:\Users\35502\Desktop\简历.docx'

# —— 个人技能替换 ——
SKILL_REPLACEMENTS = [
    (
        'AI应用开发:',
        'AI应用开发: 熟练运用 Spring AI Alibaba 框架，集成通义千问-Max、Deepseek-V3 大模型；'
        '掌握 Prompt 工程优化技巧；深入理解 RAG 原理，基于 Milvus 向量库落地稠密向量 + BM25 稀疏向量 + RRF 融合的混合检索方案；具备 AI 应用后端快速开发能力。',
    ),
    (
        'AI系统工程化:',
        'AI系统工程化: 熟悉大模型与后端系统集成方式；掌握 Tool Calling 核心技术并基于 CompletableFuture 实现多工具并行调用；'
        '了解 MCP 协议与动态客户端治理；通过 Redisson 自研 ChatMemoryRepository 实现上下文持久化；具备 AI 服务稳定性保障设计能力。',
    ),
]

# —— 新项目内容（精炼版，每条一句话）——
PROJECT_TITLE = '智策 · AI 决策助手'
PROJECT_DATE = '2026-02 至 2026-04'
TECH_STACK = 'SpringBoot + SpringCloudAlibaba(Nacos) + Spring AI Alibaba + Milvus + Redis(Redisson) + MySQL + MyBatis-Plus + RabbitMQ + MCP协议 + 通义千问-Max 等'
INTRO = '面向企业内部的 AI 决策助手平台，集成通义千问-Max/Deepseek-V3 大模型，提供基于 RAG 的知识问答、多数据源 Tool Calling 及 MCP 协议动态工具扩展能力，支持流式对话与多轮上下文记忆。'
DUTIES = '负责后端核心链路，包括 RAG 混合检索、ChatMemory 持久化、Tool Calling 工具体系、MCP 动态客户端治理及文档摄入管线的设计与实现。'

BULLETS = [
    '基于 Milvus 实现 RAG 混合检索，稠密向量与 BM25 稀疏向量双路召回并通过 RRF 融合排序，显著提升长尾查询命中率。',
    '基于 Apache Tika + TokenTextSplitter 构建文档摄入管线，区分瞬时与永久故障，结合 RabbitMQ 实现自动重试与最终一致性。',
    '基于 Redisson 自研 ChatMemoryRepository，Redis 承载热数据、RabbitMQ 异步归档 MySQL，兼顾低延迟与可追溯。',
    '基于 CompletableFuture + 自定义线程池实现 Tool Calling 并行执行，多工具 fan-out 并发触发，降低端到端延迟。',
    '自研 McpToolRegistry 实现 MCP Server 动态接入，通过 ping 探活与 SSE 自愈重连支持工具运行时热挂载和故障自动恢复。',
    '基于策略模式封装工单规则，通过 Tool Calling 让 LLM 驱动"创建-指派-通知"工单闭环，打通 AI 对话与业务落库。',
]


def set_runs_text(p_element, texts):
    """Replace texts in <w:t> children of runs in order. Only touches w:t, keeps all other attrs."""
    runs = p_element.findall(qn('w:r'))
    t_idx = 0
    for r in runs:
        for t in r.findall(qn('w:t')):
            if t_idx < len(texts):
                t.text = texts[t_idx]
                t_idx += 1
    return t_idx


def main():
    doc = docx.Document(RESUME)

    # —— 1. 更新个人技能 ——
    for prefix, new_text in SKILL_REPLACEMENTS:
        for p in doc.paragraphs:
            if p.text.startswith(prefix):
                # 清空所有 run 的文本，把新内容塞进第一个 run（保持原格式）
                runs = p._p.findall(qn('w:r'))
                for i, r in enumerate(runs):
                    for t in r.findall(qn('w:t')):
                        t.text = new_text if i == 0 else ''
                    # 第一个 run 只保留第一个 w:t
                    if i == 0:
                        break
                # 删除第一个 run 之后的所有 run
                for r in runs[1:]:
                    p._p.remove(r)
                break

    # —— 2. 定位慧云智学项目的段落模板 ——
    title_p = tech_p = intro_p = duties_p = bullet_p = anchor_p = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith('慧云智学') and title_p is None:
            title_p = p._p
        elif txt.startswith('技术栈：') and tech_p is None:
            tech_p = p._p
        elif txt.startswith('项目介绍：') and intro_p is None:
            intro_p = p._p
        elif txt.startswith('个人职责：') and duties_p is None:
            duties_p = p._p
        elif duties_p is not None and bullet_p is None and p.style.name == 'List Paragraph':
            bullet_p = p._p
        elif txt == '荣誉证书':
            anchor_p = p._p
            break

    assert all([title_p is not None, tech_p is not None, intro_p is not None,
                duties_p is not None, bullet_p is not None, anchor_p is not None]), \
        '未找到全部段落模板'

    # —— 3. 构造新项目段落（deepcopy 慧云智学结构，替换文本）——
    new_elements = []

    # 3.1 项目标题：run[0]=项目名，run[1] 包含 <w:tab/> 和 <w:t>日期
    new_title = copy.deepcopy(title_p)
    set_runs_text(new_title, [PROJECT_TITLE, PROJECT_DATE])
    new_elements.append(new_title)

    # 3.2 技术栈：run[0]=标签，run[1]=内容
    new_tech = copy.deepcopy(tech_p)
    set_runs_text(new_tech, ['技术栈：', TECH_STACK])
    new_elements.append(new_tech)

    # 3.3 项目介绍
    new_intro = copy.deepcopy(intro_p)
    set_runs_text(new_intro, ['项目介绍：', INTRO])
    new_elements.append(new_intro)

    # 3.4 个人职责
    new_duties = copy.deepcopy(duties_p)
    set_runs_text(new_duties, ['个人职责：', DUTIES])
    new_elements.append(new_duties)

    # 3.5 Bullets —— 每条 deepcopy 一次，替换文本
    for bullet_text in BULLETS:
        new_bullet = copy.deepcopy(bullet_p)
        set_runs_text(new_bullet, [bullet_text])
        new_elements.append(new_bullet)

    # —— 4. 插入到 "荣誉证书" 之前 ——
    for el in new_elements:
        anchor_p.addprevious(el)

    doc.save(RESUME)
    print('OK — rebuilt with', len(new_elements), 'paragraphs,', len(BULLETS), 'bullets')


if __name__ == '__main__':
    main()
